from flask import Blueprint, request, jsonify, send_file
from app.middleware import require_auth, get_current_user_id
from app.database import db
from app.models import Contract, Customer, User, Template, Setting, Service
from app.models.contract import contract_services
from datetime import datetime
from sqlalchemy import text
from decimal import Decimal
import os
import tempfile
from docx import Document
import re

contracts_bp = Blueprint('contracts', __name__)

@contracts_bp.route('/', methods=['GET'])
@require_auth
def get_contracts():
    """Pobiera listę kontraktów"""
    try:
        # Użyj SQL query z join żeby pobrać nazwy klientów
        contracts = db.session.execute(text("""
            SELECT c.Id, c.Title, c.ContractNumber, c.CustomerId, c.NetAmount, 
                   c.SignedAt, c.StartDate, c.EndDate, cu.Name as CustomerName
            FROM Contracts c
            LEFT JOIN Customers cu ON c.CustomerId = cu.Id
            ORDER BY c.Id DESC
        """)).fetchall()
        
        contracts_list = []
        for contract in contracts:
            contracts_list.append({
                'id': contract[0],
                'title': contract[1],
                'contractNumber': contract[2],
                'customerId': contract[3],
                'customerName': contract[8] if contract[8] else 'Brak klienta',
                'netAmount': float(contract[4]) if contract[4] else 0,
                'signedAt': contract[5].isoformat() if contract[5] else None,
                'startDate': contract[6].isoformat() if contract[6] else None,
                'endDate': contract[7].isoformat() if contract[7] else None
            })
        
        return jsonify(contracts_list), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@contracts_bp.route('/', methods=['POST'])
@require_auth
def create_contract():
    """Tworzy nowy kontrakt"""
    try:
        current_user_id = get_current_user_id()
        data = request.get_json()
        
        if not data or 'title' not in data or 'customerId' not in data:
            return jsonify({'error': 'Brak wymaganych danych'}), 400
        
        # Oblicz NetAmount na podstawie wybranych usług
        net_amount = Decimal('0')
        services_data = data.get('services', [])  # Lista obiektów: [{serviceId, quantity}]
        service_ids = data.get('serviceIds', [])  # Alternatywnie lista ID (dla kompatybilności)
        
        # Jeśli mamy serviceIds jako prostą listę, zamień na format services
        if service_ids and not services_data:
            services_data = [{'serviceId': sid, 'quantity': 1} for sid in service_ids]
        
        for service_item in services_data:
            service_id = service_item.get('serviceId') if isinstance(service_item, dict) else service_item
            quantity = service_item.get('quantity', 1) if isinstance(service_item, dict) else 1
            
            service = Service.query.get(service_id)
            if service and service.Price:
                net_amount += Decimal(str(service.Price)) * quantity
        
        # Jeśli podano netAmount ręcznie, użyj go (nadpisuje automatyczne obliczenie)
        if data.get('netAmount'):
            net_amount = Decimal(str(data.get('netAmount')))
        
        new_contract = Contract(
            Title=data['title'],
            ContractNumber=data.get('contractNumber'),
            CustomerId=data['customerId'],
            NetAmount=net_amount,
            SignedAt=datetime.fromisoformat(data['signedAt'].replace('Z', '+00:00')) if data.get('signedAt') else None,
            StartDate=datetime.fromisoformat(data['startDate'].replace('Z', '+00:00')) if data.get('startDate') else None,
            EndDate=datetime.fromisoformat(data['endDate'].replace('Z', '+00:00')) if data.get('endDate') else None,
            PaymentTermDays=data.get('paymentTermDays'),
            ScopeOfServices=data.get('scopeOfServices'),  # Zachowaj dla kompatybilności wstecznej
            CreatedByUserId=current_user_id,
            ResponsibleGroupId=data.get('responsibleGroupId')
        )
        
        db.session.add(new_contract)
        db.session.flush()  # Zapisz aby dostać contract.Id
        
        # Dodaj usługi do kontraktu
        for service_item in services_data:
            service_id = service_item.get('serviceId') if isinstance(service_item, dict) else service_item
            quantity = service_item.get('quantity', 1) if isinstance(service_item, dict) else 1
            
            service = Service.query.get(service_id)
            if service:
                # Wstaw do tabeli ContractServices używając text SQL (ponieważ mamy dodatkową kolumnę Quantity)
                db.session.execute(
                    text("INSERT INTO ContractServices (ContractId, ServiceId, Quantity) VALUES (:contract_id, :service_id, :quantity)"),
                    {'contract_id': new_contract.Id, 'service_id': service.Id, 'quantity': quantity}
                )
        
        db.session.commit()
        
        return jsonify(new_contract.to_dict()), 201
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@contracts_bp.route('/<int:contract_id>', methods=['GET'])
@require_auth
def get_contract(contract_id):
    """Pobiera pojedynczy kontrakt"""
    try:
        contract = Contract.query.get_or_404(contract_id)
        
        # Użyj metody to_dict() z modelu i dodaj customerName
        contract_data = contract.to_dict()
        contract_data['customerName'] = contract.customer.Name if contract.customer else None
        
        return jsonify(contract_data), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@contracts_bp.route('/<int:contract_id>', methods=['PUT'])
@require_auth
def update_contract(contract_id):
    """Aktualizuje kontrakt"""
    try:
        contract = Contract.query.get_or_404(contract_id)
        data = request.get_json()
        
        if not data:
            return jsonify({'error': 'Brak danych do aktualizacji'}), 400
        
        if 'title' in data:
            contract.Title = data['title']
        if 'contractNumber' in data:
            contract.ContractNumber = data['contractNumber']
        if 'customerId' in data:
            contract.CustomerId = data['customerId']
        if 'netAmount' in data:
            contract.NetAmount = data['netAmount']
        if 'signedAt' in data:
            contract.SignedAt = datetime.fromisoformat(data['signedAt'].replace('Z', '+00:00')) if data['signedAt'] else None
        if 'startDate' in data:
            contract.StartDate = datetime.fromisoformat(data['startDate'].replace('Z', '+00:00')) if data['startDate'] else None
        if 'endDate' in data:
            contract.EndDate = datetime.fromisoformat(data['endDate'].replace('Z', '+00:00')) if data['endDate'] else None
        if 'paymentTermDays' in data:
            contract.PaymentTermDays = data['paymentTermDays'] if data['paymentTermDays'] else None
        if 'scopeOfServices' in data:
            contract.ScopeOfServices = data['scopeOfServices'] if data['scopeOfServices'] else None
        
        # Aktualizuj usługi przypisane do kontraktu
        if 'services' in data or 'serviceIds' in data:
            # Usuń wszystkie istniejące powiązania z usługami
            db.session.execute(
                text("DELETE FROM ContractServices WHERE ContractId = :contract_id"),
                {'contract_id': contract.Id}
            )
            
            services_data = data.get('services', [])
            service_ids = data.get('serviceIds', [])
            
            # Jeśli mamy serviceIds jako prostą listę, zamień na format services
            if service_ids and not services_data:
                services_data = [{'serviceId': sid, 'quantity': 1} for sid in service_ids]
            
            # Oblicz nową kwotę netto
            net_amount = Decimal('0')
            for service_item in services_data:
                service_id = service_item.get('serviceId') if isinstance(service_item, dict) else service_item
                quantity = service_item.get('quantity', 1) if isinstance(service_item, dict) else 1
                
                service = Service.query.get(service_id)
                if service:
                    if service.Price:
                        net_amount += Decimal(str(service.Price)) * quantity
                    
                    # Dodaj usługę do kontraktu używając text SQL (ponieważ mamy dodatkową kolumnę Quantity)
                    db.session.execute(
                        text("INSERT INTO ContractServices (ContractId, ServiceId, Quantity) VALUES (:contract_id, :service_id, :quantity)"),
                        {'contract_id': contract.Id, 'service_id': service.Id, 'quantity': quantity}
                    )
            
            # Zaktualizuj NetAmount (chyba że podano ręcznie)
            if data.get('netAmount'):
                contract.NetAmount = Decimal(str(data.get('netAmount')))
            else:
                contract.NetAmount = net_amount
        
        if 'responsibleGroupId' in data:
            contract.ResponsibleGroupId = data['responsibleGroupId']
        
        db.session.commit()
        
        return jsonify(contract.to_dict()), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@contracts_bp.route('/<int:contract_id>', methods=['DELETE'])
@require_auth
def delete_contract(contract_id):
    """Usuwa kontrakt"""
    try:
        contract = Contract.query.get_or_404(contract_id)
        
        db.session.delete(contract)
        db.session.commit()
        
        return jsonify({'message': 'Kontrakt został usunięty'}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@contracts_bp.route('/<int:contract_id>/generate-document', methods=['GET'])
@require_auth
def generate_document(contract_id):
    """Generuje dokument kontraktu na podstawie szablonu"""
    print(f"DEBUG: generate_document called with contract_id={contract_id}")
    try:
        template_id = request.args.get('templateId')
        if not template_id:
            return jsonify({'error': 'Brak ID szablonu'}), 400
        
        # Sprawdź czy kontrakt istnieje
        contract = Contract.query.get_or_404(contract_id)
        
        # Sprawdź czy szablon istnieje
        from app.models import Template
        template = Template.query.get_or_404(template_id)
        
        # Sprawdź czy plik szablonu istnieje
        template_path = template.FilePath
        if not os.path.exists(template_path):
            return jsonify({'error': f'Plik szablonu nie istnieje: {template_path}'}), 404
        
        # Pobierz dane klienta
        customer_data = db.session.execute(text("""
            SELECT Name, Email, Phone, Company, Address, NIP, Representative
            FROM Customers 
            WHERE Id = :customer_id
        """), {'customer_id': contract.CustomerId}).fetchone()
        
        if not customer_data:
            return jsonify({'error': 'Nie znaleziono danych klienta'}), 404
        
        # Otwórz szablon Word
        doc = Document(template_path)
        
        # Przygotuj dane do zastąpienia w dokumencie
        replacements = {
            '{{CONTRACT_TITLE}}': contract.Title,
            '{{CONTRACT_NUMBER}}': contract.ContractNumber or 'Brak numeru',
            '{{CUSTOMER_NAME}}': customer_data[0] or 'Brak nazwy',
            '{{CUSTOMER_EMAIL}}': customer_data[1] or 'Brak email',
            '{{CUSTOMER_PHONE}}': customer_data[2] or 'Brak telefonu',
            '{{CUSTOMER_COMPANY}}': customer_data[3] or 'Brak firmy',
            '{{CUSTOMER_ADDRESS}}': customer_data[4] or 'Brak adresu',
            '{{CUSTOMER_NIP}}': customer_data[5] or 'Brak NIP',
            '{{CUSTOMER_REPRESENTATIVE}}': customer_data[6] or 'Brak przedstawiciela',
            '{{NET_AMOUNT}}': f"{float(contract.NetAmount):.2f}" if contract.NetAmount else "0.00",
            '{{SIGNED_DATE}}': contract.SignedAt.strftime('%d.%m.%Y') if contract.SignedAt else 'Nie podpisano',
            '{{START_DATE}}': contract.StartDate.strftime('%d.%m.%Y') if contract.StartDate else 'Nie określono',
            '{{END_DATE}}': contract.EndDate.strftime('%d.%m.%Y') if contract.EndDate else 'Nie określono',
            '{{CURRENT_DATE}}': datetime.now().strftime('%d.%m.%Y'),
        }
        
        # Zastąp placeholder'y w dokumencie
        for paragraph in doc.paragraphs:
            for placeholder, value in replacements.items():
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, str(value))
        
        # Zastąp również w tabelach
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for placeholder, value in replacements.items():
                            if placeholder in paragraph.text:
                                paragraph.text = paragraph.text.replace(placeholder, str(value))
        
        # Utwórz tymczasowy plik
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        temp_file.close()
        
        # Przygotuj nazwę pliku do pobrania
        filename = f"umowa-{contract_id}-{contract.ContractNumber or 'bez-numeru'}-{datetime.now().strftime('%Y%m%d')}.docx"
        
        # Zwróć plik do pobrania
        return send_file(
            temp_file.name,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        print(f"ERROR in generate_document: {str(e)}")
        return jsonify({'error': f'Błąd generowania dokumentu: {str(e)}'}), 500

@contracts_bp.route('/<int:contract_id>/generate-from-template', methods=['POST'])
@require_auth
def generate_contract_from_template(contract_id):
    """Generuje umowę na podstawie szablonu z wypełnieniem znaczników"""
    try:
        current_user_id = get_current_user_id()
        
        # Pobierz dane z requestu
        data = request.get_json()
        template_id = data.get('template_id')
        
        if not template_id:
            return jsonify({'error': 'ID szablonu jest wymagane'}), 400
        
        # Pobierz kontrakt
        contract = Contract.query.filter_by(Id=contract_id).first()
        if not contract:
            return jsonify({'error': 'Kontrakt nie został znaleziony'}), 404
        
        # Pobierz klienta
        customer = Customer.query.filter_by(Id=contract.CustomerId).first()
        if not customer:
            return jsonify({'error': 'Klient nie został znaleziony'}), 404
        
        # Pobierz szablon
        template = Template.query.filter_by(Id=template_id).first()
        if not template:
            return jsonify({'error': 'Szablon nie został znaleziony'}), 404
        
        # Sprawdź czy plik szablonu istnieje
        if not os.path.exists(template.FilePath):
            return jsonify({'error': 'Plik szablonu nie istnieje'}), 404
        
        # Pobierz dane firmy z bazy danych
        company_settings = {}
        settings = Setting.query.all()
        for setting in settings:
            company_settings[setting.Key] = setting.Value
        
        # Przygotuj dane do wypełnienia znaczników
        template_variables = {
            'NUMER_UMOWY': contract.ContractNumber or f"{company_settings.get('contract_prefix', 'UM')}/{datetime.now().year}/{contract.Id}",
            'DATA_PODPISANIA': contract.SignedAt.strftime('%d.%m.%Y') if contract.SignedAt else datetime.now().strftime('%d.%m.%Y'),
            'MIEJSCE_ZAWARCIA': 'Warszawa',  # Można dodać do bazy danych
            'NAZWA_KLIENTA': customer.Company or customer.Name,
            'ADRES_KLIENTA': customer.Address or 'Brak adresu',
            'NIP_KLIENTA': customer.NIP or 'Brak NIP',
            'REPREZENTANT_KLIENTA': customer.Representative or 'Brak danych',
            'NAZWA_WYKONAWCY': company_settings.get('CompanyName', 'Twoja Firma Sp. z o.o.'),
            'ADRES_WYKONAWCY': company_settings.get('CompanyAddress', 'ul. Przykładowa 123, 00-000 Warszawa'),
            'NIP_WYKONAWCY': company_settings.get('CompanyNIP', '1234567890'),
            'TYTUL_UMOWY': contract.Title or 'Świadczenie usług',
            'SZCZEGOLOWY_ZAKRES_USLUG': contract.ScopeOfServices or 'Szczegółowy zakres usług',
            'DATA_ROZPOCZECIA': contract.StartDate.strftime('%d.%m.%Y') if contract.StartDate else datetime.now().strftime('%d.%m.%Y'),
            'DATA_ZAKONCZENIA': contract.EndDate.strftime('%d.%m.%Y') if contract.EndDate else (datetime.now().replace(year=datetime.now().year + 1)).strftime('%d.%m.%Y'),
            'KWOTA_WYNAGRODZENIA_NETTO': str(int(float(contract.NetAmount))) if contract.NetAmount else '0',
            'NUMER_KONTA_BANKOWEGO_WYKONAWCY': company_settings.get('CompanyBankAccount', '12 3456 7890 1234 5678 9012 3456'),
            'TERMIN_PLATNOSCI': str(contract.PaymentTermDays) if contract.PaymentTermDays else str(company_settings.get('payment_terms_default', '14'))
        }
        
        # Wczytaj szablon
        if template.FileName.lower().endswith('.docx'):
            # Obsługa plików DOCX
            doc = Document(template.FilePath)
            
            # Zastąp znaczniki w paragrafach
            for paragraph in doc.paragraphs:
                for key, value in template_variables.items():
                    if key in paragraph.text:
                        # Znajdź wszystkie wystąpienia znacznika
                        pattern = re.compile(r'\{' + key + r'\}')
                        paragraph.text = pattern.sub(value, paragraph.text)
            
            # Zastąp znaczniki w tabelach
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for key, value in template_variables.items():
                            if key in cell.text:
                                pattern = re.compile(r'\{' + key + r'\}')
                                cell.text = pattern.sub(value, cell.text)
            
            # Zapisz do pliku tymczasowego
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
            doc.save(temp_file.name)
            
        else:
            # Obsługa plików tekstowych
            with open(template.FilePath, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Zastąp znaczniki
            for key, value in template_variables.items():
                pattern = re.compile(r'\{' + key + r'\}')
                content = pattern.sub(value, content)
            
            # Zapisz do pliku tymczasowego
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.txt', mode='w', encoding='utf-8')
            temp_file.write(content)
            temp_file.close()
        
        # Przygotuj nazwę pliku
        filename = f"umowa_{contract.ContractNumber or contract.Id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        # Zwróć plik
        return send_file(
            temp_file.name,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        print(f"ERROR in generate_contract_from_template: {str(e)}")
        return jsonify({'error': f'Błąd generowania umowy z szablonu: {str(e)}'}), 500

@contracts_bp.route('/templates', methods=['GET'])
@require_auth
def get_contract_templates():
    """Pobiera dostępne szablony umów, posortowane od najnowszych"""
    try:
        # Pobierz szablony, sortując od najnowszych (największe ID na początku)
        templates = Template.query.order_by(Template.Id.desc()).all()
        
        templates_list = []
        for template in templates:
            templates_list.append({
                'id': template.Id,
                'name': template.Name,
                'fileName': template.FileName,
                'uploadedAt': template.UploadedAt.isoformat() if template.UploadedAt else None
            })
        
        return jsonify(templates_list), 200
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500